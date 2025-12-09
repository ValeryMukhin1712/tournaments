#!/usr/bin/env python3
"""
Скрипт для экспорта Help-информации в документ Word (.docx)
"""

import re
from html.parser import HTMLParser
from html import unescape

class HelpContentExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.content = []
        self.current_tag = None
        self.current_text = ""
        
    def handle_starttag(self, tag, attrs):
        if tag in ['h4', 'h5', 'h6']:
            self.current_tag = tag
        elif tag == 'p':
            self.current_tag = 'p'
        elif tag == 'li':
            self.current_tag = 'li'
        elif tag == 'strong':
            self.current_tag = 'strong'
        elif tag == 'ul':
            self.current_tag = 'ul'
        elif tag == 'ol':
            self.current_tag = 'ol'
            
    def handle_endtag(self, tag):
        if tag in ['h4', 'h5', 'h6', 'p', 'li']:
            if self.current_text.strip():
                self.content.append({
                    'tag': self.current_tag,
                    'text': self.current_text.strip()
                })
            self.current_text = ""
            self.current_tag = None
        elif tag in ['ul', 'ol']:
            self.current_tag = None
            
    def handle_data(self, data):
        if self.current_tag:
            self.current_text += data

def extract_help_content():
    """Извлекает Help-контент из base.html"""
    with open('templates/base.html', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Находим блок help-content
    start_marker = '<div class="help-content">'
    end_marker = '</div>'
    
    start_idx = content.find(start_marker)
    if start_idx == -1:
        print("Ошибка: блок help-content не найден")
        return None
    
    # Находим закрывающий тег
    end_idx = content.find('</div>', start_idx + len(start_marker))
    if end_idx == -1:
        print("Ошибка: закрывающий тег не найден")
        return None
    
    help_html = content[start_idx + len(start_marker):end_idx]
    return help_html

def create_docx_file(help_html):
    """Создает .docx файл из HTML контента"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Парсим HTML
        parser = HelpContentExtractor()
        parser.feed(help_html)
        
        for item in parser.content:
            tag = item['tag']
            text = unescape(item['text'])
            
            if tag == 'h4':
                p = doc.add_heading(text, level=1)
                p.runs[0].font.color.rgb = RGBColor(46, 125, 50)  # Зеленый
            elif tag == 'h5':
                p = doc.add_heading(text, level=2)
                p.runs[0].font.color.rgb = RGBColor(25, 118, 210)  # Синий
            elif tag == 'h6':
                p = doc.add_heading(text, level=3)
                p.runs[0].font.color.rgb = RGBColor(66, 66, 66)  # Серый
            elif tag == 'p':
                p = doc.add_paragraph(text)
                if 'lead' in help_html:  # Проверяем класс lead
                    p.italic = True
            elif tag == 'li':
                # Убираем маркеры из текста, если они есть
                text = re.sub(r'^[•\-\*]\s*', '', text)
                doc.add_paragraph(text, style='List Bullet')
            elif tag == 'strong':
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.color.rgb = RGBColor(25, 118, 210)
        
        # Сохраняем документ
        output_file = 'help_content.docx'
        doc.save(output_file)
        print(f"✅ Документ создан: {output_file}")
        return output_file
        
    except ImportError:
        print("⚠️  Библиотека python-docx не установлена")
        print("   Установите её командой: pip install python-docx")
        print("\n   Альтернатива: используйте файл help_content.html")
        print("   Откройте его в Word и сохраните как .doc")
        return None
    except Exception as e:
        print(f"❌ Ошибка при создании .docx: {e}")
        return None

def create_rtf_file(help_html):
    """Создает RTF файл (альтернатива, если python-docx не установлен)"""
    # RTF заголовок с поддержкой Windows-1251 для кириллицы
    rtf_header = """{\\rtf1\\ansi\\ansicpg1251\\deff0\\deflang1049
{\\fonttbl{\\f0\\froman\\fprq2\\fcharset204 Times New Roman;}}
\\f0\\fs24
"""
    
    # Упрощенная конвертация HTML в текст
    text = re.sub(r'<[^>]+>', '', help_html)
    text = unescape(text)
    
    # Конвертируем текст в Windows-1251
    try:
        # Пытаемся конвертировать в Windows-1251
        text_bytes = text.encode('windows-1251', errors='replace')
        text_1251 = text_bytes.decode('windows-1251')
    except:
        # Если не получается, используем исходный текст
        text_1251 = text
    
    # Экранируем специальные символы RTF
    def escape_rtf_char(char):
        if char == '\\':
            return '\\\\'
        elif char == '{':
            return '\\{'
        elif char == '}':
            return '\\}'
        else:
            return char
    
    rtf_text = ''.join(escape_rtf_char(c) for c in text_1251)
    
    # Заменяем переносы строк на RTF команды
    rtf_text = rtf_text.replace('\n', '\\par ')
    rtf_text = rtf_text.replace('\r', '')
    
    rtf_content = rtf_header + rtf_text + "\n}"
    
    output_file = 'help_content.rtf'
    # Сохраняем в кодировке Windows-1251 для правильного отображения кириллицы
    try:
        with open(output_file, 'wb') as f:
            f.write(rtf_content.encode('windows-1251', errors='replace'))
        print(f"✅ RTF файл создан: {output_file}")
        print("   Кодировка: Windows-1251 (Cyrillic)")
        print("   Откройте его в Word и сохраните как .doc")
    except Exception as e:
        # Если не получается, сохраняем в UTF-8
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(rtf_content)
        print(f"✅ RTF файл создан: {output_file}")
        print(f"   Предупреждение: сохранено в UTF-8 (ошибка Windows-1251: {e})")
        print("   Откройте его в Word и сохраните как .doc")
    
    return output_file

if __name__ == '__main__':
    print("=" * 60)
    print("Экспорт Help-информации в документ Word")
    print("=" * 60)
    
    help_html = extract_help_content()
    if not help_html:
        exit(1)
    
    print("\n📄 Извлечен контент Help-информации")
    
    # Пытаемся создать .docx
    docx_file = create_docx_file(help_html)
    
    if not docx_file:
        # Если не получилось, создаем RTF
        print("\n📄 Создаю RTF файл как альтернативу...")
        create_rtf_file(help_html)
    
    print("\n" + "=" * 60)
    print("✅ Готово!")
    print("=" * 60)

